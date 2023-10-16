from bs4 import BeautifulSoup as BSoup
import copy

# HTML TO DOCX FILE
from htmldocx import HtmlToDocx

# def html_to_docx(html_file,docx_file):
#     new_parser = HtmlToDocx()
#     new_parser.parse_html_file(html_file, docx_file)

# Resource - https://www.e-iceblue.com/Tutorials/Python/Spire.Doc-for-Python/Program-Guide/Watermark/Python-Insert-Watermarks-in-Word.html
# Applying watermark to the document
def apply_watermark(watermark_text,font_color="grey",font_size=12):
    from spire.doc import Document, TextWatermark, WatermarkLayout, FileFormat
    from spire.doc.common import Color
    document = Document()
    document.LoadFromFile("output1.docx.docx")
    txtWatermark = TextWatermark()
    txtWatermark.Text = watermark_text
    txtWatermark.FontSize = 65
    txtWatermark.Color = Color.get_Red()
    txtWatermark.Layout = WatermarkLayout.Diagonal
    document.Watermark = txtWatermark
    document.SaveToFile("output1.docx.docx", FileFormat.Docx)
    document.Close()

# Resource - https://stackoverflow.com/questions/26752856/python-docx-set-table-cell-background-and-text-color
# Applying color to cell of table
def apply_color_to_table_cell(table_number,cell_row,cell_column,color="ff00ff"):
    print("here*********************************************************")
    from docx import Document as Doc
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    document = Doc("output1.docx.docx")
    Table = document.tables[table_number]

    #GET CELLS XML ELEMENT
    print(cell_row,cell_column)
    cell_xml_element = Table.rows[cell_row].cells[cell_column]._tc
    #RETRIEVE THE TABLE CELL PROPERTIES
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    #CREATE SHADING OBJECT
    shade_obj = OxmlElement('w:shd')
    #SET THE SHADING OBJECT

    colr = "#{0:02x}{1:02x}{2:02x}".format(max(0, min(color[0], 255)), max(0, min(color[1], 255)), max(0, min(color[2], 255)))
    shade_obj.set(qn('w:fill'), colr)
    #APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
    table_cell_properties.append(shade_obj)
    document.save("output1.docx.docx")

# Clean the HTML content
import bleach
# Load your HTML file
# Define a list of allowed HTML tags and attributes
allowed_tags = ['a', 'p', 'ul', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']
allowed_attributes = {'a': ['href', 'title']}
with open('input2.html', 'r+', encoding='utf-8') as html_file:
    html_content = html_file.read()
    html_file.truncate(0)
    cleaned_html = bleach.clean(html_content, tags=allowed_tags, attributes=allowed_attributes)
    html_file.write(cleaned_html)


# Applying bg_color_to_text
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
def apply_bgcolor_to_text(target_text, new_color):
    doc = Document("output1.docx")
    # Define the color you want to set (in RGB format)
    # color = RGBColor(new_color[0], new_color[1], new_color[2])
    for paragraph in doc.paragraphs:
        if target_text in paragraph.text:
            # Find and update the text background color by modifying the paragraph style
            paragraph.runs[0].clear()
            paragraph.runs[0].add_text(target_text)
            # paragraph.runs[0].bold = True
            paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.new_color  # Change this to the color you want

    # Save the modified document to a new file
    doc.save("output1.docx")

# Now start tree structure in python
from lxml import html
# Creating tree structure of entire html file!
def print_tree_structure(element, depth=0):
    lis.append(element.tag)
    for child in element:
        print_tree_structure(child, depth + 1)
    lis.append("/"+element.tag)

# 
lis = []
from lxml import html
def print_tree_structure(element, depth=0):
    lis.append(element.tag)
    for child in element:
        print_tree_structure(child, depth + 1)
    lis.append("/"+element.tag)

# Creating tree structure of entire html file!
def create_tree_structure(html_file):

    # html_string = '<html><body><h1>Heading 1</h1><p>Paragraph 1</p><h2>Heading 2</h2><p>Paragraph 2</p></body></html>'
    with open(html_file,'r') as html_file:
        html_string = html_file.read()
        tree = html.fromstring(html_string)
    print_tree_structure(tree)

# create_tree_structure("")

# Now parsing each tags in list and creating deque
from collections import deque 

class Object():
    bg_color = None
    text = ""
    font_size = 4
    color = None
    def __init__(self, bg_color, text,font_size,color):
        self.bg_color = bg_color,
        self.text = text,
        self.font_size = font_size,
        self.color = color


# Convert styles to inline styles
from premailer import transform
def style_inliner(html_file):
    with open(html_file,'r+') as page:
        s = transform(page.read())
        page.seek(0)
        page.truncate(0)
        page.write(s)
    # with open(html_file,'r+') as page:
    #     lines = page.readlines()
    #     page.truncate(0)
    #     page.write(lines[1:])
    return

# Dictionary to check which tag is currently open
dict = {}

def parsing_all_tags(html_file):
    style_inliner(html_file)
    new_parser = HtmlToDocx()
    new_parser.parse_html_file(html_file, "output1.docx")
    # html_to_docx(html_file,"output1.docx")
    create_tree_structure(html_file)
    open_html_file = open(html_file,"r")
    soup = BSoup(open_html_file, 'html.parser')

    row,col = -1,-1
    deq = deque()

    for i in lis:
        if(i[0]=='/'):
            temp_deque = copy.copy(deq)
            while(len(temp_deque)!=0):
                if(temp_deque[-1][1].bg_color!=None and temp_deque[-1][1].bg_color!=(None,)):
                    # deque.append_left((i,temp_deque[1]))
                    if(i=="/table"):
                        for j in range(row):
                            for k in range(col):
                                apply_color_to_table_cell(dict["table"],j,k,temp_deque[-1][1].bg_color)
                        row = 0
                        col = 0
                    if(i=="/tr"):
                        for j in range(col):
                            apply_color_to_table_cell(dict["table"],row-1,j,temp_deque[-1][1].bg_color)
                        row -= 1
                        col = 0
                    if(i=="/td" or i=="/th"):
                        print(row,col)
                        apply_color_to_table_cell(dict["table"],row,col,temp_deque[-1][1].bg_color[0])
                        col -= 1
                if(temp_deque[-1][1].color!=None):
                    apply_bgcolor_to_text(temp_deque[-1][1].text,temp_deque[-1][1].color)
                

                temp_deque.pop()
            if(len(deq)!=0):
                deq.popleft()

        else:
            # Code for bg_color,text,font_size - ToDo
            one_tag_id = None
            if(i in dict.keys()):
                one_tag_id = soup.find_all(i)[dict[i]+1].get("id")
                dict[i] += 1
            else:
                one_tag_id = soup.find_all(i)[0].get("id")
                dict[i] = 0
            all_styles = soup.find_all(i)[dict[i]].get("style")
            bg_color,text,font_size,color = None,"",4,None
            if(all_styles):
                styles_dict = {item.split(":")[0].strip(): item.split(":")[1].strip() for item in all_styles.split(";") if item}
            else:
                styles_dict = {}

            if("background-color" in styles_dict.keys()):
                bg_color = styles_dict["background-color"]
                try:
                    bg_color = bg_color[1:]
                    bg_color = tuple(int(bg_color[j:j+2], 16) for j in (0, 2, 4))
                except:
                    bg_color = None
            if("font-size" in styles_dict.keys()):
                font_size = styles_dict["font-size"]
            if(soup.find_all(i)[dict[i]].text):
                text = soup.find_all(i)[dict[i]].text.strip()
            if("color" in styles_dict.keys()):
                color = styles_dict["color"]
                print(color)
                try:
                    color = tuple(int(color[j:j+2], 16) for j in (0, 2, 4))
                except:
                    color = None

            # Check table
            # if(i=="table"):
            #     dict[]
            if(i=="tr"):
                row += 1
                col = -1
            if(i=="td" or i=="th"):
                col += 1

            print(row,col)
            # Checking watermark
            if(one_tag_id=="watermark"):
                apply_watermark(text,color,font_size)
                deq.appendleft((i,Object(None,text,font_size,None)))
            else:
                deq.appendleft((i,Object(bg_color,text,font_size,color)))


        
parsing_all_tags("D:\STGI_Hackthon\Future_scope\input.html")
